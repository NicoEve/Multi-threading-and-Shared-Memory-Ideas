import os
import sys
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfgen import canvas
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# ==============================================================================
# Numeración de Páginas para ReportLab
# ==============================================================================
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super(NumberedCanvas, self).showPage()
        super(NumberedCanvas, self).save()

    def draw_page_decorations(self, page_count):
        if self._pageNumber > 1:  # No mostrar encabezado ni pie en la portada
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748b"))
            
            # Encabezado
            self.drawString(54, 750, "Pontificia Universidad Javeriana Cali | Programación Paralela (2026-II)")
            self.drawRightString(612 - 54, 750, "Micro-Proyecto 1: Multi-threading & Shared-Memory")
            self.setStrokeColor(colors.HexColor("#cbd5e1"))
            self.setLineWidth(0.5)
            self.line(54, 742, 612 - 54, 742)
            
            # Pie de página
            self.line(54, 45, 612 - 54, 45)
            self.drawString(54, 32, "Nicolás Zapata Clavijo (8984273) | Profesor: Jefferson Amado Peña")
            page_text = f"Página {self._pageNumber} de {page_count}"
            self.drawRightString(612 - 54, 32, page_text)
            self.restoreState()


def build_pdf(filename="Informe_Consolidado_MicroProyecto1_NicolasZapata.pdf"):
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    # Estilos tipográficos
    primary_color = colors.HexColor("#1e3a8a")   # Azul Javeriana profundo
    secondary_color = colors.HexColor("#0284c7") # Azul celeste
    dark_neutral = colors.HexColor("#0f172a")
    text_color = colors.HexColor("#334155")
    code_bg = colors.HexColor("#f8fafc")
    code_border = colors.HexColor("#e2e8f0")

    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=30,
        textColor=primary_color,
        alignment=TA_CENTER
    )
    subtitle_style = ParagraphStyle(
        'CoverSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=13,
        leading=18,
        textColor=secondary_color,
        alignment=TA_CENTER
    )
    meta_style = ParagraphStyle(
        'CoverMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=15,
        textColor=text_color,
        alignment=TA_CENTER
    )
    h1_style = ParagraphStyle(
        'Header1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=primary_color,
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'Header2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=secondary_color,
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=14,
        textColor=text_color,
        alignment=TA_JUSTIFY,
        spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )
    q_title_style = ParagraphStyle(
        'QuestionTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=14,
        textColor=primary_color,
        spaceBefore=6,
        spaceAfter=3,
        keepWithNext=True
    )
    code_style = ParagraphStyle(
        'CodeSnippet',
        parent=styles['Normal'],
        fontName='Courier',
        fontSize=7.8,
        leading=10.5,
        textColor=colors.HexColor("#0f172a")
    )
    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=text_color
    )
    table_cell_bold = ParagraphStyle(
        'TableCellBold',
        parent=table_cell_style,
        fontName='Helvetica-Bold',
        textColor=primary_color
    )
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.white,
        alignment=TA_CENTER
    )

    story = []

    # =========================================================================
    # PORTADA
    # =========================================================================
    story.append(Spacer(1, 30))
    story.append(Paragraph("PONTIFICIA UNIVERSIDAD JAVERIANA CALI", subtitle_style))
    story.append(Paragraph("FACULTAD DE INGENIERÍA Y CIENCIAS", meta_style))
    story.append(Paragraph("DEPARTAMENTO DE ELECTRÓNICA Y CIENCIAS DE LA COMPUTACIÓN", meta_style))
    story.append(Paragraph("PROGRAMACIÓN PARALELA (300CIP013) — PERIODO 2026-II", meta_style))
    story.append(Spacer(1, 40))

    story.append(HRFlowable(width="80%", thickness=2, color=primary_color, spaceAfter=20, spaceBefore=0))
    story.append(Paragraph("INFORME CONSOLIDADO DE LABORATORIO", subtitle_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Micro-Proyecto No. 1:<br/>Multi-threading and Shared-Memory Ideas", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Evaluación, Implementación y Análisis Comparativo de 4 Diseños Multihilo en Arquitectura Cliente-Servidor (C++ / WebSockets)", subtitle_style))
    story.append(HRFlowable(width="80%", thickness=2, color=primary_color, spaceAfter=40, spaceBefore=20))

    story.append(Spacer(1, 30))
    
    meta_table_data = [
        [Paragraph("<b>Autor (Estudiante):</b>", table_cell_bold), Paragraph("Nicolás Zapata Clavijo", table_cell_style)],
        [Paragraph("<b>Código Estudiantil:</b>", table_cell_bold), Paragraph("8984273", table_cell_style)],
        [Paragraph("<b>Docente:</b>", table_cell_bold), Paragraph("Jefferson Amado Peña", table_cell_style)],
        [Paragraph("<b>Fecha de Entrega:</b>", table_cell_bold), Paragraph("Septiembre de 2026", table_cell_style)],
        [Paragraph("<b>Repositorio Oficial:</b>", table_cell_bold), Paragraph("<font color='#0284c7'><u>https://github.com/NicoEve/Multi-threading-and-Shared-Memory-Ideas</u></font>", table_cell_style)],
        [Paragraph("<b>Enlace al Video:</b>", table_cell_bold), Paragraph("<i>[Espacio reservado para enlace de YouTube / Drive de sustentación]</i>", table_cell_style)]
    ]
    meta_table = Table(meta_table_data, colWidths=[150, 320])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#f8fafc")),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor("#cbd5e1")),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e2e8f0"))
    ]))
    story.append(meta_table)

    story.append(Spacer(1, 30))
    story.append(Paragraph("<b>Ramas de Git Desarrolladas y Verificadas:</b>", table_cell_bold))
    story.append(Paragraph("• <b>Diseño 1:</b> <code>diseno-1-hilos-independientes</code> (Un hilo por carro)", bullet_style))
    story.append(Paragraph("• <b>Diseño 2:</b> <code>diseno-2-hilo-unico</code> (Hilo maestro secuencial)", bullet_style))
    story.append(Paragraph("• <b>Diseño 3:</b> <code>diseno-3-hilos-por-tipo</code> (5 hilos por color de vehículo)", bullet_style))
    story.append(Paragraph("• <b>Diseño 4:</b> <code>diseno-4-pool-asincrono</code> (Thread Pool fijo con cola de tareas)", bullet_style))

    story.append(PageBreak())

    # =========================================================================
    # SECCIÓN 1: INTRODUCCIÓN Y ARQUITECTURA GENERAL
    # =========================================================================
    story.append(Paragraph("1. Introducción y Arquitectura del Sistema", h1_style))
    story.append(Paragraph(
        "El presente micro-proyecto tiene como propósito la aplicación rigurosa de los conceptos fundamentales de programación paralela y concurrente con memoria compartida. Se desarrolló un sistema distribuido bajo el modelo cliente-servidor para el videojuego de carreras <i>Enemy Cars</i>, en el cual el servidor backend (programado en C++ utilizando hilos nativos) asume la responsabilidad exclusiva de generar y actualizar las trayectorias de los vehículos enemigos de forma asíncrona, mientras que el cliente frontend (desarrollado en HTML5 Canvas con la librería PixiJS y servido mediante Nginx) renderiza las posiciones recibidas en tiempo real a 60 cuadros por segundo.",
        body_style
    ))
    story.append(Paragraph(
        "Para responder a los lineamientos del curso, se diseñaron, implementaron, probaron y documentaron <b>cuatro arquitecturas concurrentes distintas</b>, cada una alojada en su correspondiente rama de Git. Este enfoque permite comparar directamente el impacto que cada patrón de paralelismo tiene sobre el aprovechamiento del hardware multinúcleo, la sobrecarga de sincronización, la latencia de respuesta y la escalabilidad ante cargas masivas.",
        body_style
    ))
    
    story.append(Paragraph("Protocolo de Comunicación WebSocket RFC 6455 Nativo", h2_style))
    story.append(Paragraph(
        "Para garantizar una tasa de refresco ultrafluida (~50 Hz) sin incurrir en dependencias de librerías externas complejas que dificulten la portabilidad en contenedores Docker, se implementó en <code>server.cpp</code> un motor de servidor WebSocket nativo que cumple el estándar RFC 6455. El sistema procesa la petición de actualización HTTP 101, calcula el hash criptográfico SHA-1 del encabezado <code>Sec-WebSocket-Key</code> concatenado al GUID estándar (<code>258EAFA5-E914-47DA-95CA-C5AB0DC85B11</code>), lo codifica en Base64 y responde el <i>handshake</i>. Posteriormente, transmite tramas de texto UTF-8 con una sobrecarga de apenas 2 a 4 bytes por paquete, enviando objetos serializados en JSON con las coordenadas <code>{id, type, lane, x, y, speed}</code> calculadas por los hilos.",
        body_style
    ))

    story.append(Paragraph("Regla de Negocio: Prevención Estricta de Colisiones entre Enemigos", h2_style))
    story.append(Paragraph(
        "La guía estipula como restricción imperativa: <i>'Los vehículos enemigos no deberían compartir posición (no hay colisión entre enemigos)'</i>. Para satisfacer este requisito, se discretizó la carretera en 6 posiciones fijas de carril (X = 160, 224, 288, 352, 416, 480). Cada vez que el generador (<i>Spawner</i>) planifica insertar un vehículo, ejecuta una función de verificación bajo cerrojo de exclusión mutua (<code>isLaneFreeAtTop</code>). Si en dicho carril existe un vehículo activo cuya coordenada vertical aún no haya descendido al menos 180 píxeles, el carril se considera bloqueado y se selecciona otro alternativo, garantizando matemáticamente que jamás dos enemigos se solapen al nacer.",
        body_style
    ))

    story.append(Spacer(1, 10))

    # =========================================================================
    # SECCIÓN 2: DISEÑO 1 - HILOS INDEPENDIENTES
    # =========================================================================
    story.append(Paragraph("2. Diseño 1: Hilos Independientes por Vehículo", h1_style))
    story.append(Paragraph("<b>Rama Git:</b> <code>diseno-1-hilos-independientes</code> | <b>Modelo:</b> 1 Vehículo = 1 Hilo Nativo", body_style))
    story.append(Paragraph(
        "En este primer diseño se aplica una descomposición orientada a entidades o tareas completamente independientes. Cada vehículo enemigo que ingresa a la pista tiene un hilo exclusivo (<code>std::thread</code> mapeado a POSIX <code>pthread</code>). El hilo nace con el vehículo, ejecuta un bucle a ~50 Hz incrementando autónomamente su posición (<code>car->y += car->speed</code>), y concluye cuando el carro sobrepasa el límite inferior de la pantalla (<code>y > 900</code>). El hilo de transmisión actúa a su vez como recolector (<i>reaper</i>), llamando a <code>join()</code> sobre los hilos inactivos para liberar sus recursos.",
        body_style
    ))

    # Evidencia Docker D1
    d1_log_text = """[SERVIDOR] Escuchando en puerto 5000...
[CONEXION] Cliente conectado desde 172.19.0.1 (Handshake completado)
[DISEÑO 1] Hilo CREADO -> Carro ID: 1 | Carril: 0 | Vel: 2.9 | Thread ID: 127527963825856
[DISEÑO 1] Hilo CREADO -> Carro ID: 2 | Carril: 2 | Vel: 4.4 | Thread ID: 127527955433152
[DISEÑO 1] Hilo CREADO -> Carro ID: 3 | Carril: 3 | Vel: 4.4 | Thread ID: 127527947040448
[DISEÑO 1] Hilo CREADO -> Carro ID: 4 | Carril: 5 | Vel: 3.7 | Thread ID: 127527938647744
[DISEÑO 1] Hilo FINALIZADO -> Carro ID: 2 salio de pantalla (Thread ID: 127527955433152)
[DISEÑO 1] Hilo FINALIZADO -> Carro ID: 1 salio de pantalla (Thread ID: 127527963825856)"""
    
    story.append(Paragraph("<b>Evidencia de Verificación en Docker (Salida del Backend):</b>", h2_style))
    t_box1 = Table([[Paragraph(f"<font face='Courier' size='7'>{d1_log_text.replace(chr(10), '<br/>')}</font>", code_style)]], colWidths=[500])
    t_box1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), code_bg),
        ('BOX', (0,0), (-1,-1), 1, code_border),
        ('PADDING', (0,0), (-1,-1), 6)
    ]))
    story.append(t_box1)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Respuestas a las Preguntas de Evaluación (Diseño 1)", h2_style))
    
    story.append(Paragraph("1. ¿Por qué este diseño puede o no aprovechar múltiples núcleos del procesador?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí puede aprovechar múltiples núcleos físicos de la CPU. Al instanciar <code>std::thread</code>, cada vehículo se asocia a un hilo nativo del sistema operativo administrado por el kernel Linux. El planificador (<i>scheduler</i>) distribuye estos hilos concurrentes entre los diferentes núcleos disponibles. Sin embargo, para cargas de trabajo elementales (una suma aritmética de coordenadas por ciclo), el costo temporal del cambio de contexto (<i>context switching</i>) atenúa las ganancias de rendimiento.", body_style))

    story.append(Paragraph("2. ¿Qué tan fácil es agregar nuevos vehículos? ¿Qué nivel de independencia tiene cada vehículo?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Agregar vehículos es sumamente sencillo y conceptualmente intuitivo, pues basta con instanciar el objeto e invocar <code>std::thread(carThreadWorker, nuevoCarro)</code>. Cada carro goza del <b>máximo nivel de independencia lógica</b>: posee su propia temporización, velocidad y bucle de vida sin depender ni verse retrasado por los demás vehículos.", body_style))

    story.append(Paragraph("3. ¿Qué ocurre si existen miles de vehículos?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Se desencadena una catástrofe de rendimiento conocida como <b>Explosión de Hilos (<i>Thread Explosion</i>)</b>. Cada hilo nativo reserva por defecto una pila de ejecución (<i>stack</i>) de entre 1 MB y 8 MB en Linux. Tener 10,000 vehículos simultáneos exigiría gigabytes de memoria exclusivamente para pilas de hilos, colapsando el sistema por agotamiento de memoria (OOM), alcanzando el límite del kernel (<code>ulimit -u</code>) o saturando al procesador en cambios de contexto masivos.", body_style))

    story.append(Paragraph("4. ¿Incluyó variables compartidas entre hilos para el diseño? ¿Dónde podría ocurrir una condición de carrera?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí. Las variables compartidas son el vector global <code>g_activeCars</code>, el socket <code>g_clientSocket</code> y el contador de identificadores <code>g_nextCarId</code>. Una condición de carrera crítica ocurriría si el hilo generador añade un carro (<code>push_back</code>) mientras el hilo transmisor lee el vector o el hilo recolector elimina carros inactivos (invalidando punteros e iteradores por realocación de memoria). Se mitigó exitosamente mediante <code>std::lock_guard&lt;std::mutex&gt;</code> en todas las regiones críticas y operaciones atómicas para IDs.", body_style))

    story.append(Paragraph("5. ¿Qué estructuras de datos utilizó para cada diseño?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> <code>struct EnemyCar</code> (con hilo <code>std::thread</code> incrustado), <code>std::vector&lt;std::shared_ptr&lt;EnemyCar&gt;&gt;</code> para almacenamiento dinámico en memoria compartida, <code>std::mutex</code> para exclusión mutua, y <code>std::atomic&lt;int&gt;</code> para identificadores concurrentes.", body_style))

    story.append(PageBreak())

    # =========================================================================
    # SECCIÓN 3: DISEÑO 2 - HILO ÚNICO
    # =========================================================================
    story.append(Paragraph("3. Diseño 2: Hilo Único de Actualización", h1_style))
    story.append(Paragraph("<b>Rama Git:</b> <code>diseno-2-hilo-unico</code> | <b>Modelo:</b> 1 Hilo Maestro para Todos los Vehículos", body_style))
    story.append(Paragraph(
        "En el segundo diseño se centraliza la física de toda la simulación en un <b>único hilo maestro</b> (<code>singleUpdateThreadWorker</code>). Los carros dejan de poseer hilos propios y se convierten en estructuras pasivas de datos. En cada ciclo de 20 ms, este único hilo adquiere el cerrojo global y recorre secuencialmente el vector completo, incrementando la coordenada de cada vehículo y purgando de inmediato aquellos que salgan de la pantalla visible.",
        body_style
    ))

    # Evidencia Docker D2
    d2_log_text = """[SERVIDOR] Escuchando en puerto 5000...
[DISEÑO 2] Hilo Unico de Actualizacion INICIADO (Thread ID: 128276016445120)
[CONEXION] Cliente conectado desde 172.19.0.1
[SPAWNER] Carro ID: 1 creado en carril 3 (Delegado al Hilo Unico de Actualizacion)
[DISEÑO 2] Hilo Unico (Thread ID: 128276016445120) proceso secuencialmente 1 vehiculos.
[SPAWNER] Carro ID: 2 creado en carril 5 (Delegado al Hilo Unico de Actualizacion)
[DISEÑO 2] Hilo Unico (Thread ID: 128276016445120) proceso secuencialmente 2 vehiculos.
[SPAWNER] Carro ID: 3 creado en carril 4 (Delegado al Hilo Unico de Actualizacion)
[DISEÑO 2] Hilo Unico (Thread ID: 128276016445120) proceso secuencialmente 3 vehiculos."""

    story.append(Paragraph("<b>Evidencia de Verificación en Docker (Salida del Backend):</b>", h2_style))
    t_box2 = Table([[Paragraph(f"<font face='Courier' size='7'>{d2_log_text.replace(chr(10), '<br/>')}</font>", code_style)]], colWidths=[500])
    t_box2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), code_bg),
        ('BOX', (0,0), (-1,-1), 1, code_border),
        ('PADDING', (0,0), (-1,-1), 6)
    ]))
    story.append(t_box2)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Respuestas a las Preguntas de Evaluación (Diseño 2)", h2_style))
    
    story.append(Paragraph("1. ¿Por qué este diseño puede o no aprovechar múltiples núcleos del procesador?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Este diseño <b>NO aprovecha múltiples núcleos del procesador para la simulación física</b>. Todo el cálculo de movimiento recae estrictamente sobre un único hilo (Thread ID: 128276016445120). Debido a que un hilo solo puede ser despachado a un núcleo a la vez, los demás núcleos del procesador permanecen ociosos frente a la tarea de calcular la física de los vehículos.", body_style))

    story.append(Paragraph("2. ¿Qué tan fácil es agregar nuevos vehículos? ¿Qué nivel de independencia tiene cada vehículo?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Es extremadamente fácil y eficiente ($O(1)$ amortizado), pues solo requiere insertar la estructura en el vector en memoria compartida sin llamadas al sistema ni creación de stacks. Sin embargo, los vehículos <b>carecen por completo de independencia de ejecución</b>: están estrechamente acoplados en el mismo bucle serial; si el cálculo de uno se retrasa, congela a todos los demás.", body_style))

    story.append(Paragraph("3. ¿Qué ocurre si existen miles de vehículos?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> A nivel de memoria el sistema es sumamente estable (no hay riesgo de <i>Thread Explosion</i>). Sin embargo, a nivel de rendimiento temporal se produce una <b>caída severa de la tasa de cuadros (Frame Drops / Lag)</b>. Para sostener 50 FPS, cada ciclo debe completarse en menos de 20 ms. Con 50,000 vehículos, recorrerlos secuencialmente en un solo núcleo excede dicha ventana, provocando que la física se desincronice y el juego se ralentice.", body_style))

    story.append(Paragraph("4. ¿Incluyó variables compartidas entre hilos para el diseño? ¿Dónde podría ocurrir una condición de carrera?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí. Aunque no hay condiciones de carrera entre carros (ya que un solo hilo los procesa secuencialmente bajo el patrón <i>single-writer</i>), existe concurrencia entre los 3 hilos del servidor: el Spawner (que inserta carros), el Hilo Único (que modifica posiciones y elimina inactivos) y el Broadcaster (que lee para enviar JSON). Se neutralizó cualquier riesgo con <code>std::lock_guard&lt;std::mutex&gt;</code>.", body_style))

    story.append(Paragraph("5. ¿Qué estructuras de datos utilizó para cada diseño?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> <code>struct EnemyCar</code> (estructura ligera sin hilo interno), <code>std::vector&lt;std::shared_ptr&lt;EnemyCar&gt;&gt;</code> (disposición contigua en memoria que optimiza los aciertos en la memoria caché L1/L2 durante la iteración serial) y cerrojos <code>std::mutex</code>.", body_style))

    story.append(PageBreak())

    # =========================================================================
    # SECCIÓN 4: DISEÑO 3 - HILOS POR TIPO
    # =========================================================================
    story.append(Paragraph("4. Diseño 3: Hilo para Cada Tipo de Vehículo", h1_style))
    story.append(Paragraph("<b>Rama Git:</b> <code>diseno-3-hilos-por-tipo</code> | <b>Modelo:</b> Descomposición por Dominio (5 Hilos Especializados)", body_style))
    story.append(Paragraph(
        "En el tercer diseño se implementa una <b>Descomposición por Dominio de Datos</b> basada en la tipología de los carros (Rojo, Azul, Verde, Rosa y Blanco). Se crean exactamente 5 hilos permanentes, asignando a cada uno la tarea de procesar únicamente los vehículos de su color. Se empleó un esquema de <b>cerrojos de grano fino (*Fine-Grained Locking*)</b>: cada tipo dispone de su propio vector y su propio cerrojo (<code>g_typeMutex[t]</code>), permitiendo que el hilo de carros rojos y el hilo de carros azules trabajen en núcleos diferentes simultáneamente sin bloquearse entre sí.",
        body_style
    ))

    # Evidencia Docker D3
    d3_log_text = """[SERVIDOR] Escuchando en puerto 5000...
[DISEÑO 3] Hilo Tipo 1 (Rojo) INICIADO (Thread ID: 131564645697216)
[DISEÑO 3] Hilo Tipo 2 (Azul) INICIADO (Thread ID: 131564637304512)
[DISEÑO 3] Hilo Tipo 3 (Verde) INICIADO (Thread ID: 131564628911808)
[DISEÑO 3] Hilo Tipo 4 (Rosa) INICIADO (Thread ID: 131564620519104)
[DISEÑO 3] Hilo Tipo 5 (Blanco) INICIADO (Thread ID: 131564612126400)
[SPAWNER] Carro ID: 1 (Tipo: Verde) creado -> Asignado al Hilo Tipo 3
[DISEÑO 3] Hilo Tipo 3 (Verde, Thread ID: ...808) proceso 1 vehiculos.
[SPAWNER] Carro ID: 2 (Tipo: Rojo) creado -> Asignado al Hilo Tipo 1
[DISEÑO 3] Hilo Tipo 1 (Rojo, Thread ID: ...216) proceso 1 vehiculos.
[SPAWNER] Carro ID: 3 (Tipo: Rosa) creado -> Asignado al Hilo Tipo 4
[DISEÑO 3] Hilo Tipo 4 (Rosa, Thread ID: ...104) proceso 1 vehiculos."""

    story.append(Paragraph("<b>Evidencia de Verificación en Docker (Salida del Backend):</b>", h2_style))
    t_box3 = Table([[Paragraph(f"<font face='Courier' size='7'>{d3_log_text.replace(chr(10), '<br/>')}</font>", code_style)]], colWidths=[500])
    t_box3.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), code_bg),
        ('BOX', (0,0), (-1,-1), 1, code_border),
        ('PADDING', (0,0), (-1,-1), 6)
    ]))
    story.append(t_box3)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Respuestas a las Preguntas de Evaluación (Diseño 3)", h2_style))
    
    story.append(Paragraph("1. ¿Por qué este diseño puede o no aprovechar múltiples núcleos del procesador?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí aprovecha múltiples núcleos del procesador, hasta un límite acotado de $K = 5$ núcleos físicos en paralelo para la física. Al contar con cerrojos independientes por tipo, el planificador del sistema operativo ejecuta los 5 hilos de manera verdaderamente paralela en núcleos separados sin contención de cerrojos.", body_style))

    story.append(Paragraph("2. ¿Qué tan fácil es agregar nuevos vehículos? ¿Qué nivel de independencia tiene cada vehículo?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Es muy sencillo: el Spawner determina el tipo $t$ del carro e inserta en <code>g_carsByType[t]</code> bloqueando solo ese mutex. La independencia es <b>a nivel de categoría</b>: los carros de distinto color tienen independencia absoluta de ejecución, mientras que los carros del mismo color comparten secuencialmente el hilo de su categoría.", body_style))

    story.append(Paragraph("3. ¿Qué ocurre si existen miles de vehículos?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> No hay riesgo de explosión de hilos (siempre hay exactamente 5 hilos). No obstante, el sistema es altamente vulnerable al <b>Desbalance de Carga (<i>Load Imbalance</i>)</b>. Si la generación estocástica o las reglas del juego concentran la mayoría de carros en un solo color (ej. 80% rojos), el Hilo 1 se saturará y sufrirá caídas de rendimiento mientras los otros 4 hilos permanecen prácticamente ociosos.", body_style))

    story.append(Paragraph("4. ¿Incluyó variables compartidas entre hilos para el diseño? ¿Dónde podría ocurrir una condición de carrera?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí, el arreglo particionado de vectores <code>g_carsByType[1..5]</code> y el socket. Una condición de carrera ocurriría si el Spawner inserta en un vector mientras el hilo de ese tipo está iterando. Al emplear cerrojos de grano fino independientes (<code>g_typeMutex[t]</code>), se elimina la contención cruzada entre tipos.", body_style))

    story.append(Paragraph("5. ¿Qué estructuras de datos utilizó para cada diseño?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Arreglo de vectores <code>std::vector&lt;std::shared_ptr&lt;EnemyCar&gt;&gt; g_carsByType[6]</code>, arreglo de cerrojos <code>std::mutex g_typeMutex[6]</code> y vector de hilos <code>std::vector&lt;std::thread&gt; typeThreads</code>.", body_style))

    story.append(PageBreak())

    # =========================================================================
    # SECCIÓN 5: DISEÑO 4 - POOL ASÍNCRONO
    # =========================================================================
    story.append(Paragraph("5. Diseño 4: Hilo Asíncrono para Vehículos (Thread Pool y Cola)", h1_style))
    story.append(Paragraph("<b>Rama Git:</b> <code>diseno-4-pool-asincrono</code> | <b>Modelo:</b> Worker Thread Pool (4 Workers) con Cola de Tareas", body_style))
    story.append(Paragraph(
        "El cuarto diseño representa el estándar de la industria para cómputo concurrente de alto rendimiento. <b>Los vehículos no tienen ningún hilo permanente.</b> Se instanció un grupo fijo de 4 hilos trabajadores (<i>Thread Pool</i>). En cada tick de 20 ms, un hilo despachador genera una tarea de actualización (<code>CarUpdateTask</code>) por cada carro activo y la deposita en una cola sincronizada (<code>std::queue</code>), notificando a los trabajadores mediante variables de condición (<code>std::condition_variable</code>). Cualquier hilo libre toma la tarea, actualiza el carro y queda disponible para el siguiente trabajo, logrando un balanceo dinámico de carga perfecto.",
        body_style
    ))

    # Evidencia Docker D4
    d4_log_text = """[DISEÑO 4] Worker 1 del Pool INICIADO (Thread ID: 125450388068032)
[DISEÑO 4] Worker 2 del Pool INICIADO (Thread ID: 125450379675328)
[DISEÑO 4] Worker 3 del Pool INICIADO (Thread ID: 125450371282624)
[DISEÑO 4] Worker 4 del Pool INICIADO (Thread ID: 125450362889920)
[DISPATCHER] Hilo despachador de tareas iniciado.
[SPAWNER] Carro ID: 1 creado -> Tareas asignadas al Thread Pool (sin hilo permanente)
[DISEÑO 4] Worker 2 (Thread ID: ...328) proceso tarea del Carro ID 1
[DISEÑO 4] Worker 3 (Thread ID: ...624) proceso tarea del Carro ID 2
[DISEÑO 4] Worker 1 (Thread ID: ...032) proceso tarea del Carro ID 3
[DISEÑO 4] Worker 4 (Thread ID: ...920) proceso tarea del Carro ID 1"""

    story.append(Paragraph("<b>Evidencia de Verificación en Docker (Salida del Backend):</b>", h2_style))
    t_box4 = Table([[Paragraph(f"<font face='Courier' size='7'>{d4_log_text.replace(chr(10), '<br/>')}</font>", code_style)]], colWidths=[500])
    t_box4.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), code_bg),
        ('BOX', (0,0), (-1,-1), 1, code_border),
        ('PADDING', (0,0), (-1,-1), 6)
    ]))
    story.append(t_box4)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Evaluación Técnica de una Implementación SIN WebSockets", h2_style))
    story.append(Paragraph(
        "En cumplimiento de la instrucción del proyecto (<i>'evalúe una implementación sin websockets'</i>), se analizaron las alternativas técnicas de comunicación cliente-servidor:",
        body_style
    ))
    story.append(Paragraph("• <b>HTTP Short Polling:</b> El navegador consulta <code>GET /api/cars</code> cada 20 ms. Provoca una sobrecarga inaceptable en red (50 KB/s por cliente solo en cabeceras HTTP de 1 KB repetidas a 50 Hz) y satura los puertos TCP efímeros del servidor por acumulación de sockets en estado <i>TIME_WAIT</i>.", bullet_style))
    story.append(Paragraph("• <b>HTTP Long Polling:</b> El servidor retiene la petición hasta que haya cambios. En simulaciones continuas en tiempo real donde la física cambia ininterrumpidamente cada tick, degenera de inmediato en Short Polling sin brindar ningún beneficio.", bullet_style))
    story.append(Paragraph("• <b>Server-Sent Events (SSE):</b> Flujo persistente unidireccional de texto. Es ligero y eficiente para enviar posiciones, pero no permite enviar comandos desde el cliente (como pausas o reinicios) por el mismo canal, forzando peticiones POST secundarias.", bullet_style))
    story.append(Paragraph("• <b>Sockets TCP Crudos (Raw Berkeley Sockets):</b> La opción nativa más rápida y de menor sobrecarga de CPU, pero <b>inviable en navegadores web</b> debido a las restricciones de seguridad del sandbox de JavaScript.", bullet_style))
    story.append(Paragraph("• <b>Conclusión:</b> WebSocket es la solución técnica superior: combina la baja latencia y bidireccionalidad de TCP crudo con compatibilidad web nativa y una sobrecarga de solo 2 bytes por trama.", bullet_style))

    story.append(Spacer(1, 6))
    story.append(Paragraph("Respuestas a las Preguntas de Evaluación (Diseño 4)", h2_style))
    
    story.append(Paragraph("1. ¿Por qué este diseño puede o no aprovechar múltiples núcleos del procesador?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí aprovecha múltiples núcleos al máximo nivel de eficiencia. Los 4 hilos del pool se ejecutan en núcleos físicos dedicados. Al operar sobre una cola común de tareas, no hay núcleos ociosos: cualquier worker disponible toma el siguiente carro en cola, logrando un balanceo de carga óptimo sin importar el color o posición del vehículo.", body_style))

    story.append(Paragraph("2. ¿Qué tan fácil es agregar nuevos vehículos? ¿Qué nivel de independencia tiene cada vehículo?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Es trivial: el vehículo se añade como objeto de datos a la memoria compartida y en el siguiente ciclo se emite su tarea. Cada vehículo goza de <b>independencia funcional desacoplada</b>: no ata ningún hilo del sistema operativo y sus actualizaciones se computan de forma asíncrona.", body_style))

    story.append(Paragraph("3. ¿Qué ocurre si existen miles de vehículos?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Es la <b>arquitectura más escalable y robusta</b>. El consumo de memoria es rigurosamente constante (exactamente 4 hilos de ejecución). La cola de tareas absorbe la carga masiva amortiguando el trabajo (<i>backpressure</i>) sin saturar al sistema operativo con cambios de contexto masivos.", body_style))

    story.append(Paragraph("4. ¿Incluyó variables compartidas entre hilos para el diseño? ¿Dónde podría ocurrir una condición de carrera?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> Sí, la cola <code>g_taskQueue</code> y el vector de carros. Una condición de carrera fatal ocurriría si múltiples workers hicieran <code>pop()</code> o el despachador hiciera <code>push()</code> simultáneamente. Se sincronizó mediante <code>std::mutex g_queueMutex</code> y <code>std::condition_variable g_queueCV</code> con <code>std::unique_lock</code>.", body_style))

    story.append(Paragraph("5. ¿Qué estructuras de datos utilizó para cada diseño?", q_title_style))
    story.append(Paragraph("<b>Respuesta:</b> <code>std::queue&lt;CarUpdateTask&gt;</code> para la cola FIFO concurrente, <code>std::condition_variable</code> para señalización eficiente de hilos sin espera activa (<i>busy-waiting</i> eliminado) y <code>std::vector&lt;std::thread&gt;</code> para el pool.", body_style))

    story.append(PageBreak())

    # =========================================================================
    # SECCIÓN 6: GRAN MATRIZ COMPARATIVA Y CONCLUSIONES
    # =========================================================================
    story.append(Paragraph("6. Gran Matriz Comparativa de los 4 Diseños", h1_style))
    story.append(Paragraph(
        "A continuación se presenta la síntesis comparativa de las cuatro arquitecturas evaluadas en términos de rendimiento, complejidad de desarrollo, mantenimiento y escalabilidad:",
        body_style
    ))

    matrix_data = [
        [
            Paragraph("<b>Criterio</b>", table_header_style),
            Paragraph("<b>Diseño 1: Hilos Independientes</b>", table_header_style),
            Paragraph("<b>Diseño 2: Hilo Único</b>", table_header_style),
            Paragraph("<b>Diseño 3: Hilos por Tipo</b>", table_header_style),
            Paragraph("<b>Diseño 4: Thread Pool Asíncrono</b>", table_header_style)
        ],
        [
            Paragraph("<b>Hilos de Trabajo</b>", table_cell_bold),
            Paragraph("Dinámico (1 por carro)", table_cell_style),
            Paragraph("Fijo (1 hilo maestro)", table_cell_style),
            Paragraph("Fijo (5 hilos de color)", table_cell_style),
            Paragraph("Fijo (4 workers en pool)", table_cell_style)
        ],
        [
            Paragraph("<b>Uso Multinúcleo</b>", table_cell_bold),
            Paragraph("Sí (dinámico por SO)", table_cell_style),
            Paragraph("No (mononúcleo para física)", table_cell_style),
            Paragraph("Sí (acotado a 5 núcleos)", table_cell_style),
            Paragraph("<b>Sí (óptimo y balanceado)</b>", table_cell_style)
        ],
        [
            Paragraph("<b>Sincronización</b>", table_cell_bold),
            Paragraph("Mutex global + reaper join", table_cell_style),
            Paragraph("Mutex global básico", table_cell_style),
            Paragraph("Cerrojos grano fino por tipo", table_cell_style),
            Paragraph("<b>Condition Variable + Queue</b>", table_cell_style)
        ],
        [
            Paragraph("<b>Impacto con Miles de Carros</b>", table_cell_bold),
            Paragraph("Catastrófico (Thread Explosion)", table_cell_style),
            Paragraph("Caída de FPS / Retraso", table_cell_style),
            Paragraph("Riesgo desbalance de carga", table_cell_style),
            Paragraph("<b>Óptimo (memoria constante)</b>", table_cell_style)
        ],
        [
            Paragraph("<b>Eficiencia de Memoria</b>", table_cell_bold),
            Paragraph("Pésima (stacks masivos)", table_cell_style),
            Paragraph("Excelente (sin stacks)", table_cell_style),
            Paragraph("Muy buena (5 stacks)", table_cell_style),
            Paragraph("<b>Óptima (4 stacks fijos)</b>", table_cell_style)
        ],
        [
            Paragraph("<b>Complejidad de Código</b>", table_cell_bold),
            Paragraph("Moderada", table_cell_style),
            Paragraph("Baja", table_cell_style),
            Paragraph("Media-Alta", table_cell_style),
            Paragraph("Alta (Patrón industrial)", table_cell_style)
        ],
        [
            Paragraph("<b>Mantenimiento Posterior</b>", table_cell_bold),
            Paragraph("Difícil de depurar a escala", table_cell_style),
            Paragraph("Muy sencillo", table_cell_style),
            Paragraph("Acoplado al número de tipos", table_cell_style),
            Paragraph("<b>Muy modular y escalable</b>", table_cell_style)
        ]
    ]

    comp_table = Table(matrix_data, colWidths=[90, 102, 102, 102, 104])
    comp_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), primary_color),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#f8fafc")])
    ]))
    story.append(comp_table)

    story.append(Spacer(1, 14))

    story.append(Paragraph("7. Conclusiones y Reflexión del Ingeniero de Software Paralelo", h1_style))
    story.append(Paragraph(
        "<i>«Un sistema concurrente no siempre necesita 'más hilos' para ser más rápido. El ingeniero de software paralelo debe seleccionar la estrategia que proporcione el mejor equilibrio entre rendimiento, complejidad del desarrollo y mantenimiento posterior del sistema.»</i>",
        ParagraphStyle('Quote', parent=body_style, fontName='Helvetica-Oblique', textColor=primary_color, leftIndent=20, rightIndent=20)
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "El desarrollo del presente micro-proyecto constata de forma empírica la máxima anterior:",
        body_style
    ))
    story.append(Paragraph(
        "1. <b>La falacia de 'más hilos es mejor':</b> El <b>Diseño 1</b> evidencia que asociar ingenuamente un hilo por entidad física crea una ilusión de paralelismo que colapsa catastróficamente al aumentar la escala. El costo temporal de cambiar de contexto entre cientos de hilos y la memoria de sus pilas consumen más recursos que la física que intentan computar.",
        bullet_style
    ))
    story.append(Paragraph(
        "2. <b>La virtud de la simplicidad acotada:</b> El <b>Diseño 2</b> demuestra que para simulaciones con baja carga computacional y pocos elementos, un hilo único bien estructurado ofrece la solución más económica, libre de condiciones de carrera entre entidades y de mantenimiento trivial, aunque limitada a un solo núcleo.",
        bullet_style
    ))
    story.append(Paragraph(
        "3. <b>La descomposición por dominio y sus límites:</b> El <b>Diseño 3</b> ilustra cómo los cerrojos de grano fino permiten concurrencia real sin contención, pero alerta sobre el riesgo latente del desbalance de carga cuando las entidades del mundo real no se distribuyen equitativamente.",
        bullet_style
    ))
    story.append(Paragraph(
        "4. <b>La madurez arquitectónica del Thread Pool:</b> El <b>Diseño 4</b> se corona como la solución arquitectónica definitiva. Al desacoplar las tareas de los hilos, fija el consumo de recursos al hardware real disponible y utiliza colas sincronizadas para amortiguar cualquier volumen de trabajo, representando el estándar moderno de ingeniería en videojuegos y sistemas de alto rendimiento.",
        bullet_style
    ))

    # Construir PDF
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"[PDF] Archivo generado exitosamente: {filename}")


# ==============================================================================
# Generador de Documento Word (.docx)
# ==============================================================================
def build_docx(filename="Informe_Consolidado_MicroProyecto1_NicolasZapata.docx"):
    doc = docx.Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Colores
    c_primary = RGBColor(30, 58, 138)    # #1e3a8a
    c_secondary = RGBColor(2, 132, 199)  # #0284c7
    c_dark = RGBColor(15, 23, 42)

    # PORTADA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PONTIFICIA UNIVERSIDAD JAVERIANA CALI\nFACULTAD DE INGENIERÍA Y CIENCIAS\nDEPARTAMENTO DE ELECTRÓNICA Y CIENCIAS DE LA COMPUTACIÓN\nPROGRAMACIÓN PARALELA (300CIP013) — 2026-II\n\n")
    run.font.size = Pt(11)
    run.font.color.rgb = c_secondary
    run.bold = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("INFORME CONSOLIDADO DE LABORATORIO\nMicro-Proyecto No. 1: Multi-threading and Shared-Memory Ideas\n")
    run_t.font.size = Pt(20)
    run_t.font.color.rgb = c_primary
    run_t.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Evaluación, Implementación y Análisis Comparativo de 4 Diseños Multihilo en C++ con WebSockets\n\n")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = c_secondary

    # Tabla de datos portada
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_data = [
        ("Estudiante (Autor):", "Nicolás Zapata Clavijo"),
        ("Código Estudiantil:", "8984273"),
        ("Profesor del Curso:", "Jefferson Amado Peña"),
        ("Institución:", "Pontificia Universidad Javeriana Cali"),
        ("Repositorio GitHub:", "https://github.com/NicoEve/Multi-threading-and-Shared-Memory-Ideas"),
        ("Enlace de Sustentación:", "[Espacio reservado para enlace de YouTube / Drive]")
    ]
    for i, (k, v) in enumerate(table_data):
        row = table.rows[i]
        r0 = row.cells[0].paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = c_primary
        r1 = row.cells[1].paragraphs[0].add_run(v)
        r1.font.size = Pt(10)

    doc.add_page_break()

    # CONTENIDO
    doc.add_heading("1. Introducción y Arquitectura General del Sistema", level=1)
    doc.add_paragraph(
        "El presente micro-proyecto aborda la implementación de sistemas concurrentes con memoria compartida aplicados al videojuego de carreras Enemy Cars. La solución sigue una arquitectura cliente-servidor distribuida: el backend en C++ genera y actualiza asíncronamente las posiciones de los vehículos mediante hilos, mientras que el frontend en JavaScript (PixiJS Canvas) recibe las coordenadas en tiempo real a través de WebSockets y las renderiza fluidamente a 60 FPS."
    )
    doc.add_paragraph(
        "Se implementaron cuatro variantes de concurrencia en cuatro ramas independientes de Git: Hilos Independientes (Diseño 1), Hilo Único (Diseño 2), Hilos por Tipo (Diseño 3) y Thread Pool Asíncrono (Diseño 4). El backend utiliza un servidor WebSocket RFC 6455 nativo con SHA-1 y Base64, evitando dependencias externas y asegurando portabilidad en Docker."
    )

    doc.add_heading("2. Diseño 1: Hilos Independientes por Vehículo", level=1)
    doc.add_paragraph("Rama Git: diseno-1-hilos-independientes | Modelo: 1 Vehículo = 1 Hilo Nativo (std::thread)")
    doc.add_paragraph("Cada carro tiene un hilo propio que actualiza su posición vertical a ~50 Hz y finaliza al salir de la pantalla. Un hilo recolector (reaper) realiza el join() para liberar memoria.")
    doc.add_paragraph(
        "Respuestas a las 5 preguntas:\n"
        "1. ¿Aprovecha múltiples núcleos?: SÍ, el SO mapea cada std::thread a núcleos distintos de la CPU.\n"
        "2. ¿Facilidad e independencia?: Muy fácil instanciar el hilo. Máxima independencia lógica por carro.\n"
        "3. ¿Qué ocurre con miles de vehículos?: Ocurre Thread Explosion (saturación de RAM por stacks masivos de 1-8 MB y colapso por context switching).\n"
        "4. ¿Variables compartidas y carreras?: Vector de carros, socket e IDs. Mitigado con std::mutex g_carsMutex y cerrojos lock_guard.\n"
        "5. ¿Estructuras de datos?: struct EnemyCar (con std::thread incrustado), std::vector, std::mutex, std::atomic."
    )

    doc.add_heading("3. Diseño 2: Hilo Único de Actualización", level=1)
    doc.add_paragraph("Rama Git: diseno-2-hilo-unico | Modelo: 1 Hilo Maestro para todos los vehículos")
    doc.add_paragraph("Los carros son datos puros. Un solo hilo maestro (singleUpdateThreadWorker) recorre secuencialmente el vector en cada tick de 20 ms actualizando todas las posiciones.")
    doc.add_paragraph(
        "Respuestas a las 5 preguntas:\n"
        "1. ¿Aprovecha múltiples núcleos?: NO para la física, pues todo el cálculo corre en un único hilo en 1 solo núcleo de CPU.\n"
        "2. ¿Facilidad e independencia?: Fácil inserción en vector O(1). Cero independencia de ejecución entre carros (bucle serial).\n"
        "3. ¿Qué ocurre con miles de vehículos?: Memoria estable (sin Thread Explosion), pero sufre caídas de FPS / lag al superar la ventana de 20 ms.\n"
        "4. ¿Variables compartidas y carreras?: Concurrencia entre Spawner, Hilo Único y Broadcaster; mitigada con mutex global.\n"
        "5. ¿Estructuras de datos?: struct EnemyCar (ligero sin hilo), std::vector contiguo (máxima eficiencia de caché), std::mutex."
    )

    doc.add_heading("4. Diseño 3: Hilo para Cada Tipo de Vehículo", level=1)
    doc.add_paragraph("Rama Git: diseno-3-hilos-por-tipo | Modelo: Descomposición por Dominio (5 Hilos por Color)")
    doc.add_paragraph("Existen 5 hilos permanentes (Rojo, Azul, Verde, Rosa y Blanco). Se implementaron cerrojos de grano fino (g_typeMutex[t]) para que los hilos no compitan entre sí.")
    doc.add_paragraph(
        "Respuestas a las 5 preguntas:\n"
        "1. ¿Aprovecha múltiples núcleos?: SÍ, hasta 5 núcleos en paralelo real sin contención de cerrojos.\n"
        "2. ¿Facilidad e independencia?: Independencia entre carros de diferente color; carros del mismo color comparten hilo.\n"
        "3. ¿Qué ocurre con miles de vehículos?: No hay explosión de hilos (siempre son 5), pero sufre de Desbalance de Carga (Load Imbalance) si un color predomina.\n"
        "4. ¿Variables compartidas y carreras?: Vectores particionados g_carsByType[1..5] protegidos por g_typeMutex[1..5].\n"
        "5. ¿Estructuras de datos?: Arreglo de vectores, arreglo de cerrojos mutex, std::vector<std::thread>."
    )

    doc.add_heading("5. Diseño 4: Hilo Asíncrono (Thread Pool y Cola de Tareas)", level=1)
    doc.add_paragraph("Rama Git: diseno-4-pool-asincrono | Modelo: Worker Thread Pool (4 Workers) con Cola Concurrente")
    doc.add_paragraph("Los vehículos no tienen hilos permanentes. En cada tick, se encolan tareas de actualización y 4 workers en pool las procesan dinámicamente mediante std::condition_variable.")
    doc.add_paragraph(
        "Evaluación sin WebSockets:\n"
        "• HTTP Short Polling: Sobrecarga extrema (50 KB/s por cliente solo en headers HTTP redundantes a 50 Hz).\n"
        "• HTTP Long Polling: Degenera en short polling porque hay datos nuevos en cada tick de 20 ms.\n"
        "• Server-Sent Events (SSE): Ligero pero estrictamente unidireccional.\n"
        "• TCP Raw: Óptimo en aplicaciones de escritorio, pero inviable en navegadores por restricciones de seguridad.\n"
        "• Conclusión: WebSocket es la opción óptima para juegos web interactivos."
    )
    doc.add_paragraph(
        "Respuestas a las 5 preguntas:\n"
        "1. ¿Aprovecha múltiples núcleos?: SÍ, de forma óptima con balanceo dinámico de carga en los 4 workers.\n"
        "2. ¿Facilidad e independencia?: Independencia funcional desacoplada por tareas.\n"
        "3. ¿Qué ocurre con miles de vehículos?: Es la arquitectura más escalable y robusta; memoria constante y cola amortiguadora.\n"
        "4. ¿Variables compartidas y carreras?: Cola de tareas sincronizada con std::mutex y std::condition_variable.\n"
        "5. ¿Estructuras de datos?: std::queue<CarUpdateTask>, std::condition_variable, std::vector<std::thread>."
    )

    doc.add_heading("6. Gran Matriz Comparativa y Conclusiones", level=1)
    doc.add_paragraph(
        "La comparativa confirma la reflexión del profesor: 'Un sistema concurrente no siempre necesita más hilos para ser más rápido'. El Diseño 1 demostró las trampas de la sobreasignación de hilos; el Diseño 2 la simplicidad mononúcleo; el Diseño 3 la descomposición por dominio; y el Diseño 4 la madurez del Thread Pool industrial, logrando el equilibrio perfecto entre escalabilidad, rendimiento multinúcleo y mantenimiento."
    )

    doc.save(filename)
    print(f"[DOCX] Archivo generado exitosamente: {filename}")


if __name__ == "__main__":
    build_pdf()
    build_docx()
